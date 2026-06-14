#!/usr/bin/env python3
"""
Generates the two main thesis tables as standalone reproducible artefacts.

Table 1 — Sample characteristics (analytic sample, N = 87).
Table 2 — Ordinal logistic regression of source attribution on text source,
          mean-centred fluency, and mean-centred credibility.

Outputs:
    ../tables/table1_sample_characteristics.csv
    ../tables/table1_sample_characteristics.docx
    ../tables/table2_ordinal_regression.csv
    ../tables/table2_ordinal_regression.docx

Dependencies:
    pandas, python-docx, statsmodels (for re-fitting), scipy

Author: Kerem Peker (ANR 2106332), Tilburg University CSAI BSc, 2026.
"""
import os
import sys
import math
import pandas as pd
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

HERE = os.path.dirname(os.path.abspath(__file__))
PROJECT_ROOT = os.path.dirname(HERE)


def _resolve_paths():
    candidates_in = [
        os.path.join(PROJECT_ROOT, 'sample_data', 'thesis_trial_data.csv'),
        os.path.join(PROJECT_ROOT, 'thesis_trial_data.csv'),
    ]
    csv_path = next((p for p in candidates_in if os.path.exists(p)), candidates_in[0])
    if os.path.basename(os.path.dirname(csv_path)) == 'sample_data':
        out_dir = os.path.join(PROJECT_ROOT, 'example_outputs', 'tables')
    else:
        out_dir = os.path.join(PROJECT_ROOT, 'tables')
    return csv_path, out_dir


INPUT_CSV, OUTPUT_DIR = _resolve_paths()
os.makedirs(OUTPUT_DIR, exist_ok=True)


# -------------------------------------------------------------------
# Helper: produce a bordered Word table within a one-page docx
# -------------------------------------------------------------------
def write_docx_table(rows: list[list[str]], outpath: str, caption: str,
                     col_widths_cm: list[float] | None = None) -> None:
    """Write a self-contained docx file with a single bordered table and caption."""
    doc = Document()
    cap = doc.add_paragraph()
    run = cap.add_run(caption)
    run.italic = True; run.font.size = Pt(10)

    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.autofit = False

    # Manual borders (the default 'Normal Table' style has none)
    tbl_pr = table._element.find(qn('w:tblPr'))
    if tbl_pr is None:
        tbl_pr = OxmlElement('w:tblPr')
        table._element.insert(0, tbl_pr)
    tbl_borders = OxmlElement('w:tblBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        bd = OxmlElement(f'w:{edge}')
        bd.set(qn('w:val'), 'single')
        bd.set(qn('w:sz'), '4')
        bd.set(qn('w:color'), '000000')
        tbl_borders.append(bd)
    tbl_pr.append(tbl_borders)

    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
                    if i == 0:
                        r.bold = True
            if col_widths_cm and j < len(col_widths_cm):
                cell.width = Cm(col_widths_cm[j])
    doc.save(outpath)


# -------------------------------------------------------------------
# Table 1 — Sample characteristics
# -------------------------------------------------------------------
def table1_sample_characteristics() -> None:
    rows = [
        ['Characteristic', 'Value'],
        ['Total responses collected', '131'],
        ['Completed (Finished = True)', '93'],
        ['Excluded (duration filter)', '6'],
        ['Analytic sample (N)', '87'],
        ['Set A / Set B', '44 / 43'],
        ['Age, M (SD)', '27.4 (9.9), range 18–63'],
        ['English proficiency (1–7), M (SD)', '4.56 (1.87)'],
        ['AI familiarity (1–7), M (SD)', '4.44 (1.96)'],
        ['Completion time, median (min)', '12.2'],
        ['Trial-level observations', '1,740'],
    ]

    pd.DataFrame(rows[1:], columns=rows[0]).to_csv(
        os.path.join(OUTPUT_DIR, 'table1_sample_characteristics.csv'),
        index=False)
    write_docx_table(
        rows,
        os.path.join(OUTPUT_DIR, 'table1_sample_characteristics.docx'),
        caption='Table 1. Sample characteristics (analytic sample, N = 87).',
        col_widths_cm=[7, 7],
    )
    print("  ✓ Table 1 → table1_sample_characteristics.{csv,docx}")


# -------------------------------------------------------------------
# Table 2 — Ordinal logistic regression (refits from raw data)
# -------------------------------------------------------------------
def table2_ordinal_regression() -> None:
    """Refit ordinal logistic regression from raw CSV and emit results."""
    if not os.path.exists(INPUT_CSV):
        sys.exit(f"ERROR: Missing {INPUT_CSV} (run _final_analysis.py first).")

    from statsmodels.miscmodels.ordinal_model import OrderedModel

    df = pd.read_csv(INPUT_CSV)
    df['fluency_c'] = df['fluency'] - df['fluency'].mean()
    df['credibility_c'] = df['credibility'] - df['credibility'].mean()

    def fmt_p(p: float) -> str:
        return '< .001' if p < 0.001 else f'.{int(round(p * 1000)):03d}'

    def fmt_beta(b: float) -> str:
        return f'{"+" if b >= 0 else "−"}{abs(b):.3f}'

    def fmt_or(b: float) -> str:
        return f'{math.exp(b):.2f}'

    def fit(predictors: list[str]):
        return OrderedModel(df['attr'], df[predictors],
                            distr='logit').fit(method='bfgs', disp=False)

    m1 = fit(['truth_AI'])
    m2 = fit(['fluency_c'])
    m3 = fit(['credibility_c'])
    m4 = fit(['truth_AI', 'fluency_c', 'credibility_c'])

    rows = [
        ['Predictor', 'β', 'SE', 'z', 'p', 'OR'],
        ['Model 1 — Truth only', '', '', '', '', ''],
        ['Truth = AI',
         fmt_beta(m1.params['truth_AI']),
         f'{m1.bse["truth_AI"]:.3f}',
         f'{m1.tvalues["truth_AI"]:+.2f}',
         fmt_p(float(m1.pvalues['truth_AI'])),
         fmt_or(m1.params['truth_AI'])],
        ['Model 2 (H4) — Fluency only', '', '', '', '', ''],
        ['Fluency (centred)',
         fmt_beta(m2.params['fluency_c']),
         f'{m2.bse["fluency_c"]:.3f}',
         f'{m2.tvalues["fluency_c"]:+.2f}',
         fmt_p(float(m2.pvalues['fluency_c'])),
         fmt_or(m2.params['fluency_c'])],
        ['Model 3 (H5) — Credibility only', '', '', '', '', ''],
        ['Credibility (centred)',
         fmt_beta(m3.params['credibility_c']),
         f'{m3.bse["credibility_c"]:.3f}',
         f'{m3.tvalues["credibility_c"]:+.2f}',
         fmt_p(float(m3.pvalues['credibility_c'])),
         fmt_or(m3.params['credibility_c'])],
        ['Model 4 — Full', '', '', '', '', ''],
        ['Truth = AI',
         fmt_beta(m4.params['truth_AI']),
         f'{m4.bse["truth_AI"]:.3f}',
         f'{m4.tvalues["truth_AI"]:+.2f}',
         fmt_p(float(m4.pvalues['truth_AI'])),
         fmt_or(m4.params['truth_AI'])],
        ['Fluency (centred)',
         fmt_beta(m4.params['fluency_c']),
         f'{m4.bse["fluency_c"]:.3f}',
         f'{m4.tvalues["fluency_c"]:+.2f}',
         fmt_p(float(m4.pvalues['fluency_c'])),
         fmt_or(m4.params['fluency_c'])],
        ['Credibility (centred)',
         fmt_beta(m4.params['credibility_c']),
         f'{m4.bse["credibility_c"]:.3f}',
         f'{m4.tvalues["credibility_c"]:+.2f}',
         fmt_p(float(m4.pvalues['credibility_c'])),
         fmt_or(m4.params['credibility_c'])],
    ]

    pd.DataFrame(rows[1:], columns=rows[0]).to_csv(
        os.path.join(OUTPUT_DIR, 'table2_ordinal_regression.csv'),
        index=False)
    write_docx_table(
        rows,
        os.path.join(OUTPUT_DIR, 'table2_ordinal_regression.docx'),
        caption=('Table 2. Ordinal logistic regression of source attribution on '
                 'text source, mean-centred fluency, and mean-centred credibility '
                 '(N = 1,740 trial-level observations across 87 participants). '
                 'β = log-odds coefficient; OR = odds ratio.'),
        col_widths_cm=[6, 2, 2, 2, 2, 2],
    )
    print("  ✓ Table 2 → table2_ordinal_regression.{csv,docx}")


def main() -> None:
    print("Generating tables...")
    table1_sample_characteristics()
    table2_ordinal_regression()
    print(f"\nAll tables saved to: {OUTPUT_DIR}")


if __name__ == '__main__':
    main()
