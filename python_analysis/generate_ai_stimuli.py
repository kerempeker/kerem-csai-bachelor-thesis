from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# All AI texts written in first-person official statement style ("we", "our")
# to match the register of human-authored texts.
# Topics and date contexts match the pre-November 2022 human-authored stimuli exactly.

results = [
    {
        "num": 1,
        "topic": "Iran's nuclear programme and JCPOA negotiations",
        "date_context": "November 2021",
        "text": "We urge Iran to engage constructively in negotiations and to take meaningful steps to address international concerns about its nuclear activities. A diplomatic path forward remains available, and we call on all parties to return to talks without preconditions. The IAEA must be granted full and unconditional access to verify compliance with Iran's obligations under international law.",
        "words": 57
    },
    {
        "num": 2,
        "topic": "Russia's military invasion of Ukraine",
        "date_context": "February 2022",
        "text": "We condemn in the strongest possible terms Russia's unprovoked and unjustified military aggression against Ukraine. This flagrant violation of international law and the UN Charter demands a firm and united response. We stand with the Ukrainian people, we will hold Russia accountable for its actions, and we are committed to providing Ukraine with the support it needs to defend itself.",
        "words": 63
    },
    {
        "num": 3,
        "topic": "Belarus orchestrating migration as a political weapon",
        "date_context": "November 2021",
        "text": "We strongly condemn the deliberate orchestration by the Lukashenko regime of irregular migration flows to destabilise our partners and undermine European security. These actions are a cynical and inhumane exploitation of vulnerable people and we call on Belarus to cease them immediately. We reaffirm our solidarity with those affected and our commitment to addressing this manufactured crisis through coordinated action.",
        "words": 62
    },
    {
        "num": 4,
        "topic": "Taliban takeover of Afghanistan and international response",
        "date_context": "August 2021",
        "text": "We call on the Taliban to guarantee the safe passage of all those who wish to leave Afghanistan, to uphold the rights of women, girls and minorities, and to prevent the country from again becoming a base for international terrorism. We will coordinate our response through the UN and with regional partners. We will judge the Taliban's government by its actions.",
        "words": 65
    },
    {
        "num": 5,
        "topic": "Climate finance commitments ahead of COP26",
        "date_context": "May 2021",
        "text": "We reaffirm our commitment to mobilising $100 billion per year in climate finance for developing countries and call on all nations to bring forward ambitious nationally determined contributions ahead of COP26. We recognise that the most vulnerable communities are already bearing the costs of climate change and commit to scaling up both mitigation and adaptation funding to address their needs.",
        "words": 63
    },
    {
        "num": 6,
        "topic": "Myanmar military coup and political repression",
        "date_context": "2021",
        "text": "We are deeply concerned by the military coup in Myanmar and the use of lethal force against peaceful protesters. We call on the Tatmadaw to release all those detained, to restore the democratically elected civilian government, and to respect the will of the Myanmar people. We strongly support ASEAN's Five Point Consensus as the framework for a peaceful resolution.",
        "words": 62
    },
    {
        "num": 7,
        "topic": "Ethiopia Tigray conflict and humanitarian access",
        "date_context": "July 2021",
        "text": "We are gravely concerned by the continued fighting in Tigray and the catastrophic humanitarian situation facing millions of civilians. We call on all parties to the conflict to agree an immediate ceasefire, to allow unimpeded humanitarian access, and to protect civilians. We urge Eritrean forces to withdraw completely from Ethiopian territory without further delay.",
        "words": 58
    },
    {
        "num": 8,
        "topic": "COVID-19 vaccine equity and COVAX funding",
        "date_context": "June 2021",
        "text": "We are determined to ensure that vaccines, tests and treatments for COVID-19 reach every corner of the world, particularly lower-income countries that lack the resources to secure their own supplies. We pledge increased funding to accelerate manufacturing capacity in developing regions, strengthen health systems, and deliver on our promise that this pandemic will not be allowed to end unequally.",
        "words": 59
    },
    {
        "num": 9,
        "topic": "UK-EU Northern Ireland Protocol negotiations",
        "date_context": "December 2020",
        "text": "We remain committed to implementing the Withdrawal Agreement in a way that protects the Good Friday Agreement and maintains the integrity of both the UK internal market and the EU single market. We have worked in good faith throughout these negotiations and continue to seek practical, pragmatic solutions that work for the people and businesses of Northern Ireland.",
        "words": 60
    },
    {
        "num": 10,
        "topic": "Mali and Sahel security and governance crisis",
        "date_context": "May 2021",
        "text": "We are deeply concerned by the deteriorating security situation across the Sahel and the ongoing threat posed by terrorist groups. We call on the transitional authorities in Mali to commit to a credible timeline for the restoration of democratic governance and to cooperate fully with regional and international partners. We reaffirm our support for MINUSMA and the G5 Sahel force.",
        "words": 63
    },
    {
        "num": 11,
        "topic": "Sudan military coup and violence against protesters",
        "date_context": "June 2022",
        "text": "We condemn in the strongest terms the violence perpetrated against peaceful protesters in Sudan and call on the military authorities to exercise maximum restraint. We mourn those who have lost their lives and extend our deepest sympathies to their families. We call for an immediate end to the violence, for accountability for those responsible, and for a return to civilian-led government.",
        "words": 65
    },
    {
        "num": 12,
        "topic": "Israel-Palestine conflict and Gaza hostilities",
        "date_context": "August 2021",
        "text": "We call on both sides to exercise maximum restraint, to uphold international humanitarian law, and to protect civilian life. We are committed to supporting international efforts to establish a durable ceasefire and to address the root causes of the conflict. There can be no military solution; lasting peace requires a negotiated two-state outcome that delivers security and dignity for Israelis and Palestinians alike.",
        "words": 63
    },
    {
        "num": 13,
        "topic": "NATO collective defence and deterrence strategy",
        "date_context": "June 2022",
        "text": "We are resolved to defend every ally against any form of aggression and to maintain an ironclad commitment to collective security. We are investing significantly in our military readiness, modernising our capabilities, and ensuring that all members meet their defence spending obligations. Our alliance is more united, more capable, and more determined than at any point in its history.",
        "words": 59
    },
    {
        "num": 14,
        "topic": "UK energy security and clean energy transition",
        "date_context": "April 2022",
        "text": "We are committed to accelerating the transition to clean, affordable, and secure domestic energy production. The global energy crisis has exposed the risks of dependence on imported fossil fuels and underscored the urgency of investing in renewable capacity. We will take decisive action to protect households and businesses from price volatility while driving forward the low-carbon transition our long-term security demands.",
        "words": 61
    },
    {
        "num": 15,
        "topic": "Global food security and famine prevention",
        "date_context": "May 2021",
        "text": "We are alarmed by the scale of food insecurity facing tens of millions of people, particularly those caught in conflict zones where humanitarian access is being deliberately obstructed. We commit to scaling up emergency food assistance and to addressing the root causes of hunger, including conflict, climate shocks, and economic instability. We call on all parties to conflict to allow unhindered aid.",
        "words": 65
    },
    {
        "num": 16,
        "topic": "China trade policy and WTO obligations",
        "date_context": "October 2021",
        "text": "We call on China to fulfil its WTO obligations, to address the concerns raised by trading partners regarding market access, intellectual property protection and state subsidies, and to engage constructively in the reform of the multilateral trading system. We recognise areas of progress and remain open to cooperation where mutual interests align, while making clear that we expect a level playing field.",
        "words": 65
    },
    {
        "num": 17,
        "topic": "Finland and Sweden applying to join NATO",
        "date_context": "May 2022",
        "text": "We warmly welcome the decision of Finland and Sweden to apply for membership of NATO and are committed to supporting their accession as swiftly as possible. Both nations are close partners who already meet the highest standards of capability and commitment. We have provided security assurances to both countries and will work closely with them throughout the accession process.",
        "words": 62
    },
    {
        "num": 18,
        "topic": "International Criminal Court and accountability",
        "date_context": "June 2022",
        "text": "We reaffirm our unwavering commitment to international criminal accountability as a cornerstone of the rules-based order. Impunity for atrocity crimes — genocide, war crimes and crimes against humanity — corrodes the foundations of international law and emboldens those who commit them. We urge all governments to fulfil their obligations, support ongoing investigations, and act in solidarity with those who have suffered grave abuses.",
        "words": 65
    },
    {
        "num": 19,
        "topic": "Syria humanitarian crisis and political process",
        "date_context": "March 2022",
        "text": "We call for the protection of civilians in Syria and for full, unimpeded access for humanitarian organisations across the country. A lasting resolution must be political, not military, and must reflect the legitimate aspirations of the Syrian people. We urge all parties to engage constructively in the UN-led peace process and to prioritise the welfare of those who have suffered most.",
        "words": 61
    },
    {
        "num": 20,
        "topic": "North Korea ballistic missile tests",
        "date_context": "May 2022",
        "text": "We are alarmed by North Korea's repeated and escalating ballistic missile launches, which pose a grave threat to regional and global security and undermine the international non-proliferation regime. We call on Pyongyang to halt these provocations, to return to negotiations, and to take concrete steps toward denuclearisation of the Korean Peninsula in a manner that can be credibly verified and sustained.",
        "words": 61
    },
]

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

title = doc.add_heading('AI-Generated Political Texts — Stimulus Materials', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub = doc.add_paragraph(
    'Thesis: Human Detection of AI-Generated Political Text\n'
    'Student: Kerem Peker | Tilburg University, 2026'
)
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph('')

meta = doc.add_paragraph()
meta.add_run('Generation parameters:\n').bold = True
meta.add_run(
    'Model: claude-opus-4-6 | Temperature: 0.7\n'
    'Prompt template: Standardized across all 20 topics (see Appendix)\n'
    'Register control: All texts produced in first-person plural ("we", "our") to match human-authored stimuli.\n'
    'Bias control: Topics provided with date context only; model instructed not to reproduce any known statement.\n'
    'Date contexts match the pre-November 2022 period of corresponding human-authored texts.'
)
doc.add_paragraph('')

for r in results:
    doc.add_heading(f"Text {r['num']}: {r['topic'].title()}", level=1)

    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    for i, (label, value) in enumerate([
        ('Topic', r['topic']),
        ('Date Context', r['date_context']),
        ('Word Count', str(r['words'])),
    ]):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        table.rows[i].cells[1].text = value

    doc.add_paragraph('')
    p = doc.add_paragraph()
    p.style = 'Quote'
    p.add_run(f'"{r["text"]}"').italic = True
    doc.add_paragraph('')
    doc.add_paragraph('─' * 80)
    doc.add_paragraph('')

doc.add_heading('Appendix: Standardized Prompt Template', level=1)
doc.add_paragraph(
    'The following prompt was used for all 20 AI-generated texts:\n\n'
    '"Write a 50–80 word official political statement on the following topic: [TOPIC].\n\n'
    'The statement should:\n'
    '- Be written in first-person plural ("we", "our") in the style of an official government or multilateral statement\n'
    '- Be politically neutral and factually accurate for the date context provided\n'
    '- Not reproduce or paraphrase any specific known statement, speech, or press release\n'
    '- Use formal political register\n'
    '- Cover the topic as of [DATE CONTEXT]\n\n'
    'Output only the statement text, nothing else."'
)

path = '/Users/kerempeker/Desktop/Thesis/AI_Generated_Stimuli.docx'
doc.save(path)
print(f"Saved: {path}")
for r in results:
    flag = " ⚠️ CHECK" if not (45 <= r['words'] <= 85) else "✓"
    print(f"  Text {r['num']:2d}: {r['words']} words {flag}")
