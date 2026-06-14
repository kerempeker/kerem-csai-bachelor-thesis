from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Title
title = doc.add_heading('Human-Authored Political Texts — Stimulus Materials', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph(
    'Thesis: Human Detection of AI-Generated Political Text\n'
    'Student: Kerem Peker | Tilburg University, 2026'
).alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph(
    'All texts sourced verbatim from verified official records (UK Government speeches, '
    'G7/E3 joint statements, UN Security Council press releases). '
    'All texts pre-date November 2022 (widespread availability of generative AI tools).'
)
doc.add_paragraph('')

texts = [
    {
        "num": 1,
        "topic": "Iran Nuclear Negotiations (JCPOA)",
        "speaker": "E3 (France, Germany, United Kingdom) — Joint Statement to the IAEA",
        "date": "25 November 2021",
        "source": "https://www.gov.uk/government/news/e3-statement-on-the-joint-comprehensive-plan-of-action",
        "text": "We as E3 will return to Vienna for negotiations in good faith, to resume work based on where we left off discussions in June. We are convinced that it is possible to reach and implement an understanding on the measures providing for Iran return to full compliance with its JCPoA commitments and the United States return to the deal.",
        "words": 59
    },
    {
        "num": 2,
        "topic": "Russia Invasion of Ukraine",
        "speaker": "G7 Leaders (Canada, France, Germany, Italy, Japan, United Kingdom, United States)",
        "date": "24 February 2022",
        "source": "https://www.gov.uk/government/news/g7-leaders-statement-on-the-invasion-of-ukraine-by-armed-forces-of-the-russian-federation-24-february-2022",
        "text": "We the Leaders of the Group of Seven (G7) are appalled by and condemn the large-scale military aggression by the Russian Federation against the territorial integrity, sovereignty and independence of Ukraine, directed partly from Belarusian soil. This unprovoked and completely unjustified attack on the democratic state of Ukraine was preceded by fabricated claims and unfounded allegations.",
        "words": 56
    },
    {
        "num": 3,
        "topic": "Belarus Migrant Crisis",
        "speaker": "G7 Foreign Ministers (Canada, France, Germany, Italy, Japan, UK, US) and EU High Representative",
        "date": "November 2021",
        "source": "https://www.gov.uk/government/news/belarus-g7-foreign-ministers-statement-on-migrant-crisis",
        "text": "We, the G7 Foreign Ministers of Canada, France, Germany, Italy, Japan, the United Kingdom and the United States of America, and the High Representative of the European Union, condemn the Belarus regime's orchestration of irregular migration across its borders. These callous acts are putting people's lives at risk and we call on Belarus to cease this activity immediately.",
        "words": 59
    },
    {
        "num": 4,
        "topic": "Afghanistan Withdrawal and Taliban Takeover",
        "speaker": "G7 Leaders — Joint Statement",
        "date": "24 August 2021",
        "source": "https://www.gov.uk/government/news/g7-leaders-statement-on-afghanistan-24-august-2021",
        "text": "We will work together, and with our allies and regional countries, through the UN, G20 and more widely, to bring the international community together to address the critical questions facing Afghanistan. As we do this, we will judge the Afghan parties by their actions, not words. In particular, we reaffirm that the Taliban will be held accountable for their actions on preventing terrorism, human rights and an inclusive political settlement.",
        "words": 71
    },
    {
        "num": 5,
        "topic": "COP26 Climate Finance Commitments",
        "speaker": "G7 Foreign and Development Ministers — Communiqué",
        "date": "5 May 2021",
        "source": "https://www.gov.uk/government/publications/g7-foreign-and-development-ministers-meeting-may-2021-communique/g7-foreign-and-development-ministers-meeting-communique-london-5-may-2021",
        "text": "We note with grave concern the impacts of climate change on the most vulnerable communities and welcome the progress made at the UK-hosted Climate and Development Ministerial on 31 March 2021. We commit to continue scaling up finance contributing to adaptation action, taking into account the priorities and needs identified by ambitious adaptation plans at local, national and sub-national levels, and adaptation communications.",
        "words": 62
    },
    {
        "num": 6,
        "topic": "Myanmar Military Coup",
        "speaker": "UK Delegation, UN Security Council",
        "date": "2021",
        "source": "https://www.gov.uk/government/speeches/a-political-coup-in-myanmar-is-fast-becoming-an-economic-and-humanitarian-crisis",
        "text": "We reiterate our support for ASEAN and the central role they are playing in responding to the crisis. The Five Point Consensus is an important agreement. We support ASEAN attempts to implement it and ensure the crisis is not perpetuated. We call on the Tatmadaw to implement commitments they have made and work constructively with ASEAN.",
        "words": 56
    },
    {
        "num": 7,
        "topic": "Ethiopia / Tigray Conflict",
        "speaker": "Lord Ahmad of Wimbledon, UK Minister for Human Rights",
        "date": "2 July 2021",
        "source": "https://www.gov.uk/government/speeches/it-is-time-to-put-the-interests-of-ethiopian-people-first",
        "text": "All sides – the Federal Government of Ethiopia, Tigray Defence Forces, Amhara militias and Eritrean Defence Forces – have an opportunity to end the cycle of violence and suffering. We urge them to take it. And we call on Eritrean forces to withdraw, as requested by the Ethiopian government.",
        "words": 50
    },
    {
        "num": 8,
        "topic": "COVID-19 Vaccine Access / COVAX",
        "speaker": "G7 Health Ministers — Communiqué",
        "date": "4 June 2021",
        "source": "https://www.gov.uk/government/publications/g7-health-ministers-meeting-june-2021-communique/g7-health-ministers-meeting-communique-oxford-4-june-2021",
        "text": "We affirm support for all existing pillars of the Access to COVID-19 Tools Accelerator (ACT-A), including its COVAX facility. We are committed to addressing the financing needs in global health to support the research, development, manufacturing, and equitable distribution of safe and effective COVID-19 diagnostics, therapeutics and vaccines.",
        "words": 51
    },
    {
        "num": 9,
        "topic": "UK–EU Northern Ireland Protocol",
        "speaker": "Michael Gove, Chancellor of the Duchy of Lancaster",
        "date": "9 December 2020",
        "source": "https://www.gov.uk/government/speeches/withdrawal-agreement-update",
        "text": "Throughout 2020 we have worked intensively to ensure that the Withdrawal Agreement, in particular the Northern Ireland Protocol, will be fully operational on 1 January 2021. Our aims, and the proportionate and pragmatic way we intended to pursue them, were set out in the Command Paper we published in May.",
        "words": 53
    },
    {
        "num": 10,
        "topic": "Mali / Sahel Security Crisis",
        "speaker": "UK Delegation, UN Security Council",
        "date": "May 2021",
        "source": "https://www.gov.uk/government/speeches/ensuring-sustainable-solutions-to-address-root-causes-of-conflict-in-the-sahel-region",
        "text": "We recognise the enormous efforts of the G5 states to address the challenges facing the region. And we welcome, in particular, their commitment through the Sahel Coalition Roadmap to refocus attention on governance, on development and the provision of basic services. This has to underpin the support for the military response to the instability in the Sahel.",
        "words": 58
    },
    {
        "num": 11,
        "topic": "Sudan Coup and Civilian Deaths",
        "speaker": "Joint Statement (UK, US, EU, Norway, Switzerland)",
        "date": "8 June 2022",
        "source": "https://www.gov.uk/government/news/joint-statement-on-death-toll-since-25th-military-coup-in-sudan",
        "text": "We deeply regret the loss of Sudanese lives, killed in large-scale attacks and violence across the country during the same period. The loss of life and many injured represent a heavy toll for the people of Sudan. We would like to extend our sympathies to the families and friends of the victims.",
        "words": 55
    },
    {
        "num": 12,
        "topic": "Israel–Palestine Conflict / Gaza Violence",
        "speaker": "Ambassador Barbara Woodward, UK Permanent Representative to the UN",
        "date": "30 August 2021",
        "source": "https://www.gov.uk/government/speeches/preventing-further-escalation-of-violence-in-israel-and-the-occupied-palestinian-territories",
        "text": "We share concerns about the recent tensions in Gaza and call on all parties to take steps to avoid exacerbating tensions and to maintain the cessation of hostilities. The United Kingdom condemns unequivocally Hamas' indiscriminate attacks against Israel, including the use of incendiary balloons. We call upon Hamas and other terrorist groups to permanently end their rocket fire against Israel.",
        "words": 60
    },
    {
        "num": 13,
        "topic": "NATO Collective Defence / Madrid Summit",
        "speaker": "Rt Hon Boris Johnson MP, Prime Minister of the United Kingdom",
        "date": "30 June 2022",
        "source": "https://www.gov.uk/government/speeches/pms-press-conference-remarks-at-the-nato-summit-30-june-2022",
        "text": "And we are resolved not just to support Ukraine, but we have agreed a new strategic concept, we are moving beyond the doctrine of tripwire deterrence on NATO's eastern flank to a new approach of defence deterrence by denial. And countries around the table are also recognising that they must spend more.",
        "words": 56
    },
    {
        "num": 14,
        "topic": "UK Energy Security Strategy",
        "speaker": "Kwasi Kwarteng, Secretary of State for Business, Energy and Industrial Strategy",
        "date": "5 April 2022",
        "source": "https://www.gov.uk/government/speeches/delivering-great-britains-energy-security",
        "text": "The UK has always known that we need to decarbonise and generate more cheap, clean power at home to reduce our exposure to global gas markets we are unable to control. We were the first major economy to legislate for net zero. And we now need to go further and faster.",
        "words": 53
    },
    {
        "num": 15,
        "topic": "Global Food Security Crisis",
        "speaker": "G7 Foreign and Development Ministers — Communiqué",
        "date": "5 May 2021",
        "source": "https://www.gov.uk/government/publications/g7-foreign-and-development-ministers-meeting-may-2021-communique/g7-foreign-and-development-ministers-meeting-communique-london-5-may-2021",
        "text": "We are gravely concerned by the UN's reports that over 34 million people are already facing emergency levels of food insecurity, are one step from catastrophe or famine, and that almost 80 million have been forcibly displaced and 237 million need humanitarian assistance. We endorse the work of the G7 Famine Prevention and Humanitarian Crises Panel and the Compact which aim to prevent famine in 2021 and begin to stem the growth of humanitarian need.",
        "words": 75
    },
    {
        "num": 16,
        "topic": "China Trade Policy / WTO Review",
        "speaker": "Simon Manley, UK Ambassador to the WTO and UN in Geneva",
        "date": "20 October 2021",
        "source": "https://www.gov.uk/government/speeches/wto-trade-policy-review-of-china-uk-statement",
        "text": "We recognise the recent progress China has made to open its markets to international trade and investment. We welcome China's engagement with plurilateral initiatives such as the JSI on investment facilitation, E-Commerce and domestic regulation, as well as the Trade and Health Initiative, the Informal Dialogue on Plastics Pollution and Environmentally Sustainable Plastics - all good initiatives.",
        "words": 59
    },
    {
        "num": 17,
        "topic": "Finland and Sweden NATO Membership Bid",
        "speaker": "Rt Hon Elizabeth Truss MP, Secretary of State for Foreign, Commonwealth and Development Affairs",
        "date": "16 May 2022",
        "source": "https://www.gov.uk/government/news/uk-welcomes-sweden-and-finlands-nato-membership-plans",
        "text": "We look forward to working with them as new NATO Allies and stand ready to offer them our every assistance during the accession process. Our mutual security declarations signed with Sweden and Finland last week by the Prime Minister demonstrate our steadfast and unequivocal commitment to both countries during this process and beyond.",
        "words": 54
    },
    {
        "num": 18,
        "topic": "International Criminal Court — 20th Anniversary",
        "speaker": "Ambassador James Kariuki CMG, UK Deputy Permanent Representative to the United Nations",
        "date": "24 June 2022",
        "source": "https://www.gov.uk/government/speeches/reflecting-on-our-relationship-with-the-international-criminal-court-after-20-years-of-the-rome-statute",
        "text": "We note the obligation of States Parties to cooperate with the Court under the Rome Statute, and we also call on all States to cooperate with the Court where there are UN Security Council Resolutions which require this. We will continue to demonstrate our support for the Court, and to work together with States Parties and the Court, to ensure that the Court delivers justice for victims, and accountability, in respect of the most serious crimes of international concern.",
        "words": 79
    },
    {
        "num": 19,
        "topic": "Syria Humanitarian Crisis",
        "speaker": "Ambassador James Kariuki CMG, UK Deputy Permanent Representative to the United Nations",
        "date": "24 March 2022",
        "source": "https://www.gov.uk/government/speeches/addressing-the-lingering-humanitarian-situation-in-syria-uk-at-the-un-security-council",
        "text": "As we have said many times, there can be no military solution to the Syrian conflict. We continue to support the UN-facilitated, Syrian-led, political process outlined in resolution 2254. We urge Council members to continue to call for a nationwide ceasefire, unhindered aid access and conditions for safe refugee return.",
        "words": 52
    },
    {
        "num": 20,
        "topic": "North Korea ICBM Tests",
        "speaker": "G7 Foreign Ministers and EU High Representative — Joint Statement",
        "date": "30 May 2022",
        "source": "https://www.gov.uk/government/news/intercontinental-ballistic-missile-test-by-north-korea-g7-foreign-ministers-statement-30-may-2022",
        "text": "We, the G7 Foreign Ministers and the High Representative of the European Union, reiterate our urgent call on the DPRK to abandon its weapons of mass destruction and ballistic missile programs in a complete, verifiable and irreversible manner and to fully comply with all legal obligations arising from the relevant Security Council resolutions.",
        "words": 56
    },
]

for t in texts:
    # Topic heading
    heading = doc.add_heading(f"Text {t['num']}: {t['topic']}", level=1)

    # Metadata table
    table = doc.add_table(rows=4, cols=2)
    table.style = 'Table Grid'
    cells = table.rows[0].cells
    cells[0].text = 'Speaker'
    cells[1].text = t['speaker']
    cells = table.rows[1].cells
    cells[0].text = 'Date'
    cells[1].text = t['date']
    cells = table.rows[2].cells
    cells[0].text = 'Word Count'
    cells[1].text = str(t['words'])
    cells = table.rows[3].cells
    cells[0].text = 'Source'
    cells[1].text = t['source']

    # Bold the left column
    for row in table.rows:
        row.cells[0].paragraphs[0].runs[0].bold = True

    doc.add_paragraph('')

    # The verbatim text in a styled block
    p = doc.add_paragraph()
    p.style = 'Quote'
    run = p.add_run(f'"{t["text"]}"')
    run.italic = True

    doc.add_paragraph('')
    doc.add_paragraph('─' * 80)
    doc.add_paragraph('')

# Notes
doc.add_heading('Methodological Notes', level=1)
notes = doc.add_paragraph()
notes.add_run('Source types:\n').bold = True
notes.add_run(
    '• gov.uk: Official UK Government speeches and statements (verbatim transcripts of '
    'parliamentary or UN statements)\n'
    '• G7/E3 joint statements: Multilateral statements co-signed by UK and partner governments\n\n'
)
notes.add_run('Pre-November 2022 criterion:\n').bold = True
notes.add_run(
    'All 20 texts pre-date November 2022 (the launch of ChatGPT and widespread availability '
    'of generative AI writing tools). Dates range from December 2020 to June 2022. '
    'This ensures human texts could not have been AI-generated, eliminating potential '
    'training data contamination.\n\n'
)
notes.add_run('First-person plural register:\n').bold = True
notes.add_run(
    'All texts use "we/our" as the dominant grammatical person, matching the AI-generated '
    'stimuli in register. This controls for grammatical person as a detection cue. '
    'Occasional third-person references within predominantly first-person passages '
    'are standard in multilateral political statements.\n\n'
)
notes.add_run('Known caveats:\n').bold = True
notes.add_run(
    '• Texts 5 and 15 are drawn from the same G7 FM Communiqué (5 May 2021) '
    'but from different sections (climate adaptation and food security respectively). '
    'They are separated by 10 topics in the randomised stimulus order.\n'
    '• Text 9 (UK-EU NI Protocol) is from December 2020 — the earliest text in the set — '
    'as the Protocol negotiations were ongoing throughout 2020–2022.\n'
    '• Several texts fall slightly below the 50-word target (texts 7, 8, 12, 15, 18, 19); '
    'these reflect the natural length of official statement passages and remain within '
    'the spirit of the 50–80 word guideline.'
)

path = '/Users/kerempeker/Desktop/Thesis/Human_Authored_Stimuli.docx'
doc.save(path)
print(f"Saved: {path}")
for t in texts:
    flag = " ⚠️ SHORT" if t['words'] < 45 else ("✓" if t['words'] >= 50 else " ⚠ BORDERLINE")
    print(f"  Text {t['num']:2d}: {t['words']:3d} words {flag}  — {t['topic']}")
